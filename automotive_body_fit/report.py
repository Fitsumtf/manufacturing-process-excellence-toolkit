"""Word report generator for the automotive body-fit case study."""

from pathlib import Path

import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .analysis import CapabilityResult

DISCLAIMER = (
    "This project uses a synthetic educational dataset created solely to demonstrate "
    "statistical process capability analysis. It does not contain proprietary or "
    "confidential manufacturing data from any company."
)

REFERENCES = [
    ("AIAG Statistical Process Control (SPC), 2nd Edition", "https://www.aiag.org/training-and-resources/manuals/details/SPC-3"),
    ("AIAG Measurement Systems Analysis (MSA), 4th Edition", "https://www.aiag.org/training-and-resources/manuals/details/MSA-4"),
    ("AIAG & VDA FMEA Handbook", "https://www.aiag.org/training-and-resources/manuals"),
    ("Douglas C. Montgomery, Introduction to Statistical Quality Control, 8th Edition", "https://www.wiley.com/en-us/Introduction%2Bto%2BStatistical%2BQuality%2BControl%2C%2B8th%2BEdition-p-9781119399308"),
    ("ISO 22514 series — Statistical methods in process management", "https://www.iso.org/committee/49742/x/catalogue/"),
]


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _hyperlink(paragraph, label: str, url: str) -> None:
    rel_id = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single")
    props.extend([color, underline]); run.append(props)
    text = OxmlElement("w:t"); text.text = label
    run.append(text); hyperlink.append(run); paragraph._p.append(hyperlink)


def _metric_table(doc: Document, result: CapabilityResult) -> None:
    rows = [
        ("Samples", f"{result.n}"), ("Mean", f"{result.mean:.6f} mm"),
        ("Sample STDEV", f"{result.sample_stdev:.6f} mm"),
        ("Specification", f"{result.lsl:.3f} to {result.usl:.3f} mm"),
        ("Cp", f"{result.cp:.3f}"), ("Cpk", f"{result.cpk:.3f}"),
        ("Observed rejects", f"{result.observed_rejects} ({result.observed_ppm:,.0f} ppm)"),
        ("Expected rejects*", f"{result.expected_ppm_normal:,.0f} ppm"),
    ]
    table = doc.add_table(rows=4, cols=4)
    table.style = "Light Shading Accent 1"
    for idx, (label, value) in enumerate(rows):
        r, c = divmod(idx, 2)
        c *= 2
        table.cell(r, c).text = label
        table.cell(r, c + 1).text = value
        _shade(table.cell(r, c), "D9EAF2")
        for cell in (table.cell(r, c), table.cell(r, c + 1)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build_report(
    values: np.ndarray, result: CapabilityResult, histogram: Path, run_chart: Path,
    output: Path, *, report_title: str = "Automotive Body Fit Process Capability Study",
    context: str = "High-Volume Automotive OEM | Hood-to-Fender Gap | Synthetic Educational Case Study",
    characteristic: str = "hood-to-fender gap",
    disclaimer: str = DISCLAIMER,
) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7); section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8); section.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"; styles["Normal"].font.size = Pt(10)
    for name in ("Title", "Heading 1", "Heading 2"):
        styles[name].font.name = "Aptos Display"; styles[name].font.color.rgb = RGBColor(23, 50, 77)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(report_title)
    subtitle = doc.add_paragraph(context)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note = doc.add_paragraph()
    note.style = doc.styles["Intense Quote"]
    note.add_run("PORTFOLIO DISCLAIMER — ").bold = True
    note.add_run(disclaimer)
    _metric_table(doc, result)

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        f"The study evaluates {result.n} measurements for {characteristic}. "
        f"The mean is {result.mean:.3f} against a {result.target:.3f} target. "
        f"The sample standard deviation is {result.sample_stdev:.3f} mm. Cp is {result.cp:.3f} and Cpk is {result.cpk:.3f}, "
        f"with an illustrative status of '{result.status}'. Cp/Cpk interpretation is valid only after "
        "measurement-system adequacy and statistical stability have been established."
    )
    p = doc.add_paragraph(); p.add_run(f"Conclusion: {result.status}. ").bold = True
    if result.cp < 1.0:
        p.add_run("The estimated natural process spread is wider than the available specification window.")
    elif result.cpk < result.cp * 0.95:
        p.add_run("The process location materially reduces capability relative to its potential capability.")
    else:
        p.add_run("The process is reasonably centered; confirm stability and the applicable acceptance criterion.")

    doc.add_heading("2. Definitions and Equations", level=1)
    doc.add_paragraph("Cp = (USL - LSL) / (6 x s)")
    doc.add_paragraph("Cpk = min[(USL - mean) / (3 x s), (mean - LSL) / (3 x s)]")
    doc.add_paragraph(
        f"For this dataset, Cp = ({result.usl:.3f} - {result.lsl:.3f}) / "
        f"(6 x {result.sample_stdev:.3f}) = {result.cp:.3f}. Because the mean is centered, "
        f"Cpu = Cpl = Cpk = {result.cpk:.3f}. Calculations use the sample standard deviation (n - 1)."
    )

    doc.add_heading("3. Visual Analysis", level=1)
    for image, caption in ((histogram, "Figure 1. Capability histogram and fitted normal distribution."), (run_chart, "Figure 2. Time-ordered measurements and specification limits.")):
        doc.add_picture(str(image), width=Inches(6.2))
        p = doc.add_paragraph(caption); p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("4. Engineering Interpretation", level=1)
    points = [
        "Centering is not the primary problem: the mean matches the nominal target.",
        f"Variation is the primary problem: six standard deviations equal {result.six_sigma_spread:.3f} mm, wider than the {result.spec_width:.3f} mm specification window.",
        f"Observed out-of-specification count is {result.observed_rejects}, equivalent to {result.observed_ppm:,.0f} ppm in this dataset.",
        "The normal-model expected PPM is an estimate, not a guarantee of future defect rate.",
        "Confirm statistical stability with an appropriate control chart and measurement adequacy with Gauge R&R before relying on capability indices.",
    ]
    for item in points: doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5. Recommended Improvement Plan", level=1)
    actions = [
        ("Measurement", "Perform Gauge R&R and standardize gauge orientation and contact pressure."),
        ("Equipment/fixture", "Inspect locators, clamps, tooling interfaces, and wear points; establish preventive-maintenance limits."),
        ("Method", "Standardize adjustment sequence, torque method, and acceptance criteria in the work instruction."),
        ("Material/design", "Review tolerance-stack contributors, process inputs, material variation, and supplier variation."),
        ("Control", "Use rational subgroups and an X-bar/R chart, or an I-MR chart when subgrouping is not appropriate."),
        ("Verification", "After corrective action, collect an independent time-ordered dataset and demonstrate the program-specific Cpk requirement."),
    ]
    table = doc.add_table(rows=1, cols=2); table.style = "Light Shading Accent 1"
    table.rows[0].cells[0].text = "Workstream"; table.rows[0].cells[1].text = "Recommended action"
    for cell in table.rows[0].cells: _shade(cell, "17324D")
    for workstream, action in actions:
        cells = table.add_row().cells; cells[0].text = workstream; cells[1].text = action

    doc.add_heading("6. Statistical Cautions", level=1)
    doc.add_paragraph(
        "*Expected PPM assumes an independent, stable, approximately normal process. Cp/Cpk do not establish stability, "
        "normality, causal control, or measurement-system adequacy. Specification limits are engineering requirements, "
        "not control limits. Acceptance thresholds must follow applicable customer and program requirements."
    )

    doc.add_heading("7. References", level=1)
    for label, url in REFERENCES:
        p = doc.add_paragraph(style="List Bullet"); _hyperlink(p, label, url)

    doc.add_heading("Appendix A — Raw Measurements", level=1)
    table = doc.add_table(rows=1, cols=3); table.style = "Light Shading Accent 1"
    for cell, label in zip(table.rows[0].cells, ("Sample", "Value (mm)", "Status")):
        cell.text = label; _shade(cell, "17324D")
    for i, value in enumerate(values, start=1):
        cells = table.add_row().cells
        cells[0].text = str(i); cells[1].text = f"{value:.6f}"
        cells[2].text = "PASS" if result.lsl <= value <= result.usl else "FAIL"

    footer = section.footer.paragraphs[0]
    footer.text = "Interactive Process Capability Analyzer | Engineering review required"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output
