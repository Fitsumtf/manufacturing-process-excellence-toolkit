"""Generate a concise all-in-one manufacturing analytics portfolio report."""

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from automotive_body_fit.analysis import analyze_capability, load_measurements
from .gage_rr import crossed_gage_rr
from .oee import calculate_oee
from .pfmea import analyze_pfmea
from .process_data import pareto_summary

DISCLAIMER = (
    "This portfolio uses synthetic educational datasets and generalized manufacturing "
    "case studies solely to demonstrate engineering, statistical analysis, and problem-solving "
    "methods. It contains no proprietary production records, confidential specifications, "
    "internal documents, or company-specific manufacturing data."
)


def _heading(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def _kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2); table.style = "Light Shading Accent 1"
    table.rows[0].cells[0].text = "Metric"; table.rows[0].cells[1].text = "Result"
    for key, value in rows:
        cells = table.add_row().cells; cells[0].text = key; cells[1].text = value


def build_toolkit_report(data_dir: Path, output: Path) -> Path:
    values = load_measurements(data_dir / "automotive_body_fit_100_samples.csv")
    cap = analyze_capability(values, lsl=3.0, target=3.5, usl=4.0)
    pfmea = analyze_pfmea(pd.read_csv(data_dir / "synthetic_pfmea.csv"))
    grr = crossed_gage_rr(pd.read_csv(data_dir / "synthetic_gage_rr_data.csv"))
    oee_data = pd.read_csv(data_dir / "synthetic_oee_data.csv")
    oee_results = [calculate_oee(**row) for row in oee_data.drop(columns="date").to_dict("records")]
    defects = pd.read_csv(data_dir / "synthetic_defect_data.csv")
    pareto = pareto_summary(defects, "defect_category", "count")

    doc = Document(); section = doc.sections[0]
    section.top_margin = Inches(.7); section.bottom_margin = Inches(.7)
    section.left_margin = Inches(.8); section.right_margin = Inches(.8)
    doc.styles["Normal"].font.name = "Aptos"; doc.styles["Normal"].font.size = Pt(10)
    for style in ("Title", "Heading 1", "Heading 2"):
        doc.styles[style].font.color.rgb = RGBColor(23, 50, 77)
    title = doc.add_paragraph(style="Title"); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Manufacturing Process Excellence Toolkit")
    subtitle = doc.add_paragraph("Integrated Python Portfolio: Capability | PFMEA | OEE | Gauge R&R | SPC | Process Improvement")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(); p.add_run("PUBLIC PORTFOLIO DISCLAIMER — ").bold = True; p.add_run(DISCLAIMER)

    _heading(doc, "1. Executive Summary")
    doc.add_paragraph(
        "This report demonstrates an integrated manufacturing-engineering workflow. Capability evaluates output "
        "variation; PFMEA prioritizes preventive action; OEE quantifies production losses; Gauge R&R evaluates "
        "measurement adequacy; process-data tools expose instability and dominant losses; generalized case studies "
        "show how these methods support hands-on line ownership and technical problem solving."
    )

    _heading(doc, "2. Process Capability")
    _kv_table(doc, [("Samples", str(cap.n)), ("Mean", f"{cap.mean:.6f} mm"),
                    ("Sample standard deviation", f"{cap.sample_stdev:.6f} mm"),
                    ("Cp / Cpk", f"{cap.cp:.3f} / {cap.cpk:.3f}"),
                    ("Observed nonconformance", f"{cap.observed_rejects} ({cap.observed_ppm:,.0f} ppm)"),
                    ("Interpretation", cap.status)])

    _heading(doc, "3. PFMEA Risk Review")
    doc.add_paragraph(
        "RPN is calculated as Severity x Occurrence x Detection. Portfolio priority is a transparent educational "
        "triage and is not the proprietary AIAG-VDA Action Priority table."
    )
    table = doc.add_table(rows=1, cols=4); table.style = "Light Shading Accent 1"
    for c, h in zip(table.rows[0].cells, ["Process step", "Failure mode", "RPN", "Revised RPN"]): c.text = h
    for _, row in pfmea.iterrows():
        cells = table.add_row().cells
        for cell, value in zip(cells, [row.process_step, row.failure_mode, str(row.rpn), str(row.revised_rpn)]): cell.text = value

    _heading(doc, "4. OEE Trend")
    table = doc.add_table(rows=1, cols=5); table.style = "Light Shading Accent 1"
    for c, h in zip(table.rows[0].cells, ["Date", "Availability", "Performance", "Quality", "OEE"]): c.text = h
    for date, result in zip(oee_data.date, oee_results):
        cells = table.add_row().cells
        for cell, value in zip(cells, [date, f"{result.availability:.1%}", f"{result.performance:.1%}", f"{result.quality:.1%}", f"{result.oee:.1%}"]): cell.text = value

    _heading(doc, "5. Gauge R&R")
    _kv_table(doc, [("Study design", f"{grr.parts} parts x {grr.operators} operators x {grr.trials} trials"),
                    ("Total Gauge R&R (% study variation)", f"{grr.percent_study_variation_grr:.2f}%"),
                    ("Repeatability", f"{grr.percent_study_variation_repeatability:.2f}%"),
                    ("Reproducibility", f"{grr.percent_study_variation_reproducibility:.2f}%"),
                    ("Number of distinct categories", str(grr.ndc)),
                    ("Interpretation", grr.interpretation)])

    _heading(doc, "6. Process Data and Pareto")
    table = doc.add_table(rows=1, cols=4); table.style = "Light Shading Accent 1"
    for c, h in zip(table.rows[0].cells, ["Defect category", "Count", "Percent", "Cumulative"]): c.text = h
    for _, row in pareto.iterrows():
        cells = table.add_row().cells
        for cell, value in zip(cells, [row.defect_category, f"{row['count']:.0f}", f"{row.percent:.1f}%", f"{row.cumulative_percent:.1f}%"]): cell.text = value

    _heading(doc, "7. Hands-On Engineering and Problem Solving")
    for title, text in [
        ("Process launch", "Converted an uncontrolled repair need into defined work content, tooling, measurement, standardized work, training, and production controls."),
        ("Line redesign", "Used line observation, cycle-time analysis, task precedence, work balancing, layout, controls, and ramp support to improve flow and expose bottlenecks."),
        ("Vision commissioning", "Supported installation, recipe development, correlation, false-call reduction, reaction plans, and ongoing verification with an equipment supplier."),
        ("Variation reduction", "Connected MSA, time-ordered SPC, fixture inspection, standardized method, tolerance review, corrective action, and independent capability validation."),
    ]:
        p = doc.add_paragraph(style="List Bullet"); p.add_run(f"{title}: ").bold = True; p.add_run(text)

    _heading(doc, "8. Professional Use and Limitations")
    doc.add_paragraph(
        "Results are educational demonstrations. Production decisions require approved specifications, rational "
        "sampling, stable-process evidence, validated measurement systems, customer-specific acceptance criteria, "
        "and review by qualified process and quality personnel."
    )
    output.parent.mkdir(parents=True, exist_ok=True); doc.save(output); return output
