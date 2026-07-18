"""Downloadable Word report for Version 2 live production intelligence."""

from __future__ import annotations

from io import BytesIO

import pandas as pd
from docx import Document


def build_v2_report(records: pd.DataFrame, metrics: dict, query: str, matches: pd.DataFrame, brief: str) -> bytes:
    doc = Document()
    doc.add_heading("Manufacturing Process Intelligence Report", 0)
    doc.add_paragraph(
        "Synthetic educational portfolio — no proprietary or confidential manufacturing data."
    )
    doc.add_heading("Executive KPIs", level=1)
    table = doc.add_table(rows=1, cols=2); table.style = "Light Shading Accent 1"
    table.rows[0].cells[0].text = "Metric"; table.rows[0].cells[1].text = "Result"
    rows = [
        ("Production records", str(metrics.get("records", 0))),
        ("Units started", str(metrics.get("units_started", 0))),
        ("First Pass Yield", f"{metrics.get('fpy', 0):.2%}"),
        ("Rolled Throughput Yield", f"{metrics.get('rty', 0):.2%}"),
        ("Weighted OEE", f"{metrics.get('weighted_oee', 0):.2%}"),
        ("Reworked / scrapped", f"{metrics.get('reworked', 0)} / {metrics.get('scrapped', 0)}"),
    ]
    for key, value in rows:
        cells = table.add_row().cells; cells[0].text = key; cells[1].text = value
    doc.add_heading("Production Records", level=1)
    doc.add_paragraph(records.to_string(index=False))
    doc.add_heading("Problem Search", level=1); doc.add_paragraph(query or "No query entered")
    if not matches.empty:
        for row in matches.itertuples():
            doc.add_heading(f"{row.lesson_id}: {row.problem}", level=2)
            doc.add_paragraph(f"BM25 score: {row.bm25_score:.3f}")
            doc.add_paragraph(f"Root cause: {row.root_cause}")
            doc.add_paragraph(f"Corrective action: {row.corrective_action}")
            doc.add_paragraph(f"Verification: {row.verification}")
    doc.add_heading("Engineering Brief", level=1); doc.add_paragraph(brief)
    stream = BytesIO(); doc.save(stream); return stream.getvalue()
